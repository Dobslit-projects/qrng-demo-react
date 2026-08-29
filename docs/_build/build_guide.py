#!/usr/bin/env python3
"""Converte docs/GUIA_DO_USUARIO_KAPUA.md em HTML (para PDF via Playwright) e DOCX
(python-docx). Parser de Markdown restrito ao subconjunto usado no guia:
headings #..######, parágrafos, listas '- ', tabelas GFM '|', code fences ```,
**negrito**, `código`, [texto](url), blockquote '>', hr '---', <div align=...>.

Uso: python docs/_build/build_guide.py
Saídas: docs/GUIA_DO_USUARIO_KAPUA.html  e  docs/GUIA_DO_USUARIO_KAPUA.docx
"""
from __future__ import annotations
import html
import pathlib
import re
import sys

ROOT = pathlib.Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "GUIA_DO_USUARIO_KAPUA.md"
OUT_HTML = ROOT / "docs" / "GUIA_DO_USUARIO_KAPUA.html"
OUT_DOCX = ROOT / "docs" / "GUIA_DO_USUARIO_KAPUA.docx"

INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*(?!\*)([^*\n]+?)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def split_cells(row: str):
    """Divide uma linha de tabela GFM em células, respeitando `código` (não
    divide em '|' dentro de crases) e o escape '\\|'."""
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    cells, buf, in_code, k = [], [], False, 0
    while k < len(row):
        c = row[k]
        if c == "\\" and k + 1 < len(row) and row[k + 1] == "|":
            buf.append("|"); k += 2; continue
        if c == "`":
            in_code = not in_code; buf.append(c); k += 1; continue
        if c == "|" and not in_code:
            cells.append("".join(buf).strip()); buf = []; k += 1; continue
        buf.append(c); k += 1
    cells.append("".join(buf).strip())
    return cells


def parse_blocks(text: str):
    """Gera uma lista de blocos: ('h',lvl,txt) ('p',txt) ('ul',[items]) ('table',[rows])
    ('code',txt) ('hr',) ('center',txt) ('quote',[lines])."""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        # HTML comment
        if s.startswith("<!--"):
            while i < n and "-->" not in lines[i]:
                i += 1
            i += 1
            continue
        # centered div
        if s.startswith("<div align"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("</div>"):
                t = lines[i].strip()
                if t:
                    m2 = re.match(r"(#{1,6})\s+(.*)", t)
                    if m2:
                        buf.append((f"h{len(m2.group(1))}", m2.group(2).strip()))
                    else:
                        buf.append(("p", t))
                i += 1
            i += 1
            blocks.append(("center", buf))
            continue
        # code fence
        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(buf)))
            continue
        # hr
        if s == "---":
            blocks.append(("hr",))
            i += 1
            continue
        # heading
        m = re.match(r"(#{1,6})\s+(.*)", s)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        # blockquote
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append(("quote", buf))
            continue
        # table
        if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?\s*$", lines[i + 1].strip()):
            rows = [split_cells(s)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_cells(lines[i]))
                i += 1
            blocks.append(("table", rows))
            continue
        # list
        if re.match(r"^[-*]\s+", s):
            items = []
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ul", items))
            continue
        # ordered list
        if re.match(r"^\d+\.\s+", s):
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ol", items))
            continue
        # paragraph (join consecutive non-empty non-special lines)
        buf = [s]
        i += 1
        while i < n:
            t = lines[i].strip()
            if (not t or t.startswith(("#", "|", ">", "```", "- ", "* ", "---", "<div"))
                    or re.match(r"^\d+\.\s+", t)):
                break
            buf.append(t)
            i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks


# ---------- HTML ----------
def inline_html(t: str) -> str:
    t = html.escape(t)
    t = INLINE_CODE.sub(lambda m: f"<code>{m.group(1)}</code>", t)
    t = BOLD.sub(lambda m: f"<strong>{m.group(1)}</strong>", t)
    t = ITALIC.sub(lambda m: f"<em>{m.group(1)}</em>", t)
    t = LINK.sub(lambda m: f'<a href="{m.group(2)}">{m.group(1)}</a>', t)
    return t


def to_html(blocks) -> str:
    out = []
    for b in blocks:
        k = b[0]
        if k == "h":
            out.append(f"<h{b[1]}>{inline_html(b[2])}</h{b[1]}>")
        elif k == "p":
            out.append(f"<p>{inline_html(b[1])}</p>")
        elif k == "center":
            parts = []
            for kind, txt in b[1]:
                if kind.startswith("h"):
                    parts.append(f"<{kind}>{inline_html(txt)}</{kind}>")
                else:
                    parts.append(f"<p>{inline_html(txt)}</p>")
            cls = "center cover" if any(kk == "h1" for kk, _ in b[1]) else "center"
            out.append(f'<div class="{cls}">{"".join(parts)}</div>')
        elif k == "hr":
            out.append("<hr>")
        elif k == "code":
            out.append(f"<pre><code>{html.escape(b[1])}</code></pre>")
        elif k == "quote":
            out.append("<blockquote>" + "<br>".join(inline_html(x) for x in b[1]) + "</blockquote>")
        elif k == "ul":
            out.append("<ul>" + "".join(f"<li>{inline_html(x)}</li>" for x in b[1]) + "</ul>")
        elif k == "ol":
            out.append("<ol>" + "".join(f"<li>{inline_html(x)}</li>" for x in b[1]) + "</ol>")
        elif k == "table":
            rows = b[1]
            thead = "<tr>" + "".join(f"<th>{inline_html(c)}</th>" for c in rows[0]) + "</tr>"
            body = "".join(
                "<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in r) + "</tr>"
                for r in rows[1:]
            )
            out.append(f"<table><thead>{thead}</thead><tbody>{body}</tbody></table>")
    css = """
body{font:13px/1.5 -apple-system,Segoe UI,Roboto,Arial,sans-serif;color:#1a1a1a;max-width:920px;margin:0 auto;padding:32px}
h1{font-size:26px;border-bottom:3px solid #0c8ce9;padding-bottom:6px;margin-top:0}
h2{font-size:20px;border-bottom:1px solid #ccc;padding-bottom:4px;margin-top:28px}
h3{font-size:16px;margin-top:20px}
h4{font-size:14px;margin-top:16px}
code{background:#f2f4f7;padding:1px 4px;border-radius:3px;font-family:ui-monospace,Consolas,monospace;font-size:12px}
pre{background:#0f1419;color:#e6edf3;padding:12px 14px;border-radius:6px;overflow-x:auto}
pre code{background:none;color:inherit;padding:0;font-size:12px}
table{border-collapse:collapse;width:100%;margin:12px 0;font-size:11.5px}
th,td{border:1px solid #d0d7de;padding:5px 8px;text-align:left;vertical-align:top}
th{background:#f2f4f7;font-weight:700}
tr:nth-child(even) td{background:#fafbfc}
blockquote{border-left:4px solid #0c8ce9;background:#f0f8ff;margin:10px 0;padding:8px 14px;color:#333}
hr{border:none;border-top:1px solid #ddd;margin:22px 0}
.center{text-align:center;margin:18px 0}
.cover{margin:32mm 0;page-break-after:always}
.cover h1{font-size:40px;border:none;margin-bottom:4px}
.cover h2{font-size:20px;border:none;color:#0c8ce9;margin-top:0}
.cover p{color:#444;margin:4px 0}
a{color:#0c56b3;word-break:break-all}
table{page-break-inside:auto}
tr{page-break-inside:avoid}
h2,h3,h4{page-break-after:avoid}
pre{page-break-inside:avoid}
@page{margin:18mm 16mm}
"""
    return f"<!doctype html><meta charset='utf-8'><title>Kapuã — Guia do Usuário</title><style>{css}</style>\n" + "\n".join(out)


# ---------- DOCX ----------
def to_docx(blocks, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def add_inline(paragraph, text):
        # handle **bold** and `code`; links become "texto (url)"
        text = LINK.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*(?!\*)[^*\n]+?\*(?!\*))", text)
        for p in parts:
            if not p:
                continue
            if p.startswith("**") and p.endswith("**"):
                r = paragraph.add_run(p[2:-2]); r.bold = True
            elif p.startswith("`") and p.endswith("`"):
                r = paragraph.add_run(p[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
            elif len(p) > 1 and p.startswith("*") and p.endswith("*"):
                r = paragraph.add_run(p[1:-1]); r.italic = True
            else:
                paragraph.add_run(p)

    for b in blocks:
        k = b[0]
        if k == "h":
            doc.add_heading(b[2].replace("`", "").replace("**", ""), level=min(b[1], 4))
        elif k == "p":
            add_inline(doc.add_paragraph(), b[1])
        elif k == "center":
            is_cover = any(kk == "h1" for kk, _ in b[1])
            for kind, txt in b[1]:
                if kind == "h1":
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(txt); r.bold = True; r.font.size = Pt(28)
                elif kind.startswith("h"):
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(txt); r.bold = True; r.font.size = Pt(16)
                else:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(p, txt)
            if is_cover:
                doc.add_page_break()
        elif k == "hr":
            doc.add_paragraph("—" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif k == "code":
            p = doc.add_paragraph()
            r = p.add_run(b[1]); r.font.name = "Consolas"; r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Pt(10)
        elif k == "quote":
            p = doc.add_paragraph()
            r = p.add_run("  ".join(b[1])); r.italic = True
            p.paragraph_format.left_indent = Pt(14)
        elif k == "ul":
            for it in b[1]:
                add_inline(doc.add_paragraph(style="List Bullet"), it)
        elif k == "ol":
            for it in b[1]:
                add_inline(doc.add_paragraph(style="List Number"), it)
        elif k == "table":
            rows = b[1]
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=len(rows), cols=ncol)
            t.style = "Light Grid Accent 1"
            for ri, r in enumerate(rows):
                for ci in range(ncol):
                    cell = t.rows[ri].cells[ci]
                    cell.text = ""
                    add_inline(cell.paragraphs[0], r[ci] if ci < len(r) else "")
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)
                        if ri == 0:
                            run.bold = True
    doc.save(str(path))


def main():
    md = SRC.read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    OUT_HTML.write_text(to_html(blocks), encoding="utf-8")
    print(f"HTML -> {OUT_HTML}  ({OUT_HTML.stat().st_size} bytes)")
    try:
        to_docx(blocks, OUT_DOCX)
        print(f"DOCX -> {OUT_DOCX}  ({OUT_DOCX.stat().st_size} bytes)")
    except Exception as e:  # noqa: BLE001
        print(f"DOCX FALHOU: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
